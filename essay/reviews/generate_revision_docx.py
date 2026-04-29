"""
论文修订对比Word文档生成器 (Revision DOCX Generator for Thesis)

作者: Claude (Auto-generated)
版本: 1.0.0
日期: 2026-04-24
位置: essay/reviews/generate_revision_docx.py

功能说明:
    将原始markdown与修订后的markdown进行对比，生成带有红色标记的Word文档。
    红色字体表示新增或修改的内容，黑色字体表示原文保留未改的内容。
    便于导师/审稿人快速定位修改位置。

核心逻辑:
    1. 读取原始稿与修订稿两份markdown文件
    2. 提取原始稿中所有段落文本作为"基准指纹"
    3. 遍历修订稿，逐段判断该段落是否在原始稿中存在
       - 若不存在（或相似度低于阈值），判定为新增/修改内容，标红
       - 若用户显式指定了red_sections（如"2.2.2,2.2.3"），该章节下所有内容强制标红
    4. 生成Word文档，支持markdown粗体(**text**)解析

字体与格式:
    - 正文字体: 仿宋 (中文) / Times New Roman (英文), 小四号(12pt)
    - 标题字体: 仿宋, 三号(16pt), 加粗
    - 代码字体: Courier New, 10pt, 红色

依赖:
    pip install python-docx

用法示例 (命令行):
    # 生成第二章修订对比文档
    python essay/reviews/generate_revision_docx.py \
        --original essay/chapters/chapter_02_draft.md \
        --final essay/reviews/Chapter_02/CHAPTER_02_FINAL.md \
        --output essay/output/chapter_02_revised_red.docx \
        --title "第二章 相关技术与理论基础（修订稿）" \
        --red-sections "2.2.2,2.2.3,2.3.1,2.3.2,2.3.3,2.4.1,2.4.2,2.4.3" \
        --legend "红色=新增/修改；黑色=原文保留。主要修改：(1)2.2.2改为分词+知识图谱；(2)2.2.3扩写LLM规范语义；..."

用法示例 (Python模块导入):
    from generate_revision_docx import generate_revision_docx

    generate_revision_docx(
        original_md_path="essay/chapters/chapter_02_draft.md",
        final_md_path="essay/reviews/Chapter_02/CHAPTER_02_FINAL.md",
        output_docx_path="essay/output/chapter_02_revised_red.docx",
        title="第二章 相关技术与理论基础（修订稿）",
        red_sections={"2.2.2", "2.2.3", "2.3.1", "2.3.2", "2.3.3", "2.4.1", "2.4.2", "2.4.3"},
        legend_text="红色=新增/修改；黑色=原文保留。"
    )

注意事项:
    1. red_sections 参数用于标记"整节重写"的情况（如2.2.2从OpenRouter改为分词+知识图谱）
    2. 对于"扩写"而非"重写"的章节，可以不填red_sections，脚本会自动逐段比对标红新增段落
    3. 段落比对基于前60个字符的指纹（去除空格），阈值默认为65%重叠度
    4. 生成后的Word文档建议在Microsoft Word中打开查看，WPS可能渲染略有差异
"""

import argparse
import re
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn


RED = RGBColor(255, 0, 0)
BLACK = RGBColor(0, 0, 0)
FONT_NAME = "仿宋"
FONT_NAME_EN = "Times New Roman"
FONT_SIZE = Pt(12)
FONT_SIZE_TITLE = Pt(16)
FONT_SIZE_CODE = Pt(10)
CODE_FONT = "Courier New"


def _set_run_font(run, font_name=FONT_NAME, font_size=FONT_SIZE, color=BLACK, bold=False):
    """统一设置run的字体属性"""
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if font_size:
        run.font.size = font_size
    run.font.color.rgb = color
    run.font.bold = bold


def _extract_original_paragraphs(original_md: str) -> set:
    """提取原始markdown中的段落，用于后续对比判断是否为新增内容"""
    paras = set()
    for line in original_md.split("\n"):
        line = line.strip()
        if not line or line.startswith("#") or line.startswith("---") or line.startswith("```"):
            continue
        # 取前60个字符作为指纹（去除空格）
        key = line.replace(" ", "").replace("　", "")[:60]
        if len(key) > 10:
            paras.add(key)
    return paras


def _is_new_paragraph(text: str, original_paras: set, min_overlap_ratio: float = 0.65) -> bool:
    """判断一段文本是否在原始稿中不存在（即为新增/修改内容）"""
    key = text.replace(" ", "").replace("　", "")[:60]
    if key in original_paras:
        return False
    # 检查是否有高重叠度的原始段落
    for op in original_paras:
        if key in op or op in key:
            return False
        if len(key) > 30 and len(op) > 30:
            matches = sum(1 for a, b in zip(key, op) if a == b)
            ratio = matches / max(len(key), len(op))
            if ratio > min_overlap_ratio:
                return False
    return True


def _parse_markdown(final_md: str):
    """
    解析markdown，生成结构化块列表。
    每个块为字典：{"type": "header"|"paragraph"|"code", "content": str, "level": int(仅header)}
    """
    blocks = []
    lines = final_md.split("\n")
    in_code = False
    code_buffer = []
    para_buffer = []

    def flush_para():
        if para_buffer:
            text = " ".join(para_buffer)
            blocks.append({"type": "paragraph", "content": text})
            para_buffer.clear()

    def flush_code():
        if code_buffer:
            for code_line in code_buffer:
                blocks.append({"type": "code", "content": code_line})
            code_buffer.clear()

    for raw_line in lines:
        stripped = raw_line.strip()

        if stripped.startswith("```"):
            flush_para()
            if in_code:
                flush_code()
            in_code = not in_code
            continue

        if in_code:
            para_buffer.clear()  # 确保段落缓冲已清空
            code_buffer.append(stripped)
            continue

        if stripped.startswith("#"):
            flush_para()
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped.lstrip("#").strip()
            blocks.append({"type": "header", "content": text, "level": level})
            continue

        if not stripped:
            flush_para()
            continue

        para_buffer.append(stripped)

    flush_para()
    return blocks


def _write_paragraph(doc, text: str, is_red: bool, bold_pattern: str = r"\*\*(.*?)\*\*"):
    """向文档写入一个段落，支持markdown粗体标记，支持红色高亮"""
    p = doc.add_paragraph()
    parts = re.split(bold_pattern, text)
    color = RED if is_red else BLACK
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        is_bold = (i % 2 == 1)  # **text** 中的 text 为粗体
        _set_run_font(run, color=color, bold=is_bold)
    return p


def generate_revision_docx(
    original_md_path: str,
    final_md_path: str,
    output_docx_path: str,
    title: str,
    red_sections: set = None,
    legend_text: str = None,
):
    """
    核心函数：生成带红色修订标记的Word文档。

    参数：
        original_md_path: 原始markdown文件路径
        final_md_path: 修订后markdown文件路径
        output_docx_path: 输出Word文档路径
        title: 文档标题（如"第二章 相关技术与理论基础（修订稿）"）
        red_sections: 需要标红的章节编号集合，如{"2.2.2", "2.2.3"}
        legend_text: 修改说明文字，显示在文档开头
    """
    red_sections = red_sections or set()

    with open(original_md_path, "r", encoding="utf-8") as f:
        original_md = f.read()
    with open(final_md_path, "r", encoding="utf-8") as f:
        final_md = f.read()

    original_paras = _extract_original_paragraphs(original_md)
    blocks = _parse_markdown(final_md)

    doc = Document()

    # 设置默认样式
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME_EN
    style.font.size = FONT_SIZE
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)

    # 标题
    title_para = doc.add_heading(title, level=0)
    title_para.alignment = 1  # 居中
    for run in title_para.runs:
        _set_run_font(run, font_size=FONT_SIZE_TITLE, bold=True)

    # 修改说明
    if legend_text:
        legend_para = doc.add_paragraph()
        run_label = legend_para.add_run("【修改说明】")
        _set_run_font(run_label, color=RED, bold=True)
        run_desc = legend_para.add_run(" " + legend_text)
        _set_run_font(run_desc)
        doc.add_paragraph()

    current_section = ""

    for block in blocks:
        btype = block["type"]

        if btype == "header":
            text = block["content"]
            level = block["level"]
            p = doc.add_heading(text, level=level)
            for run in p.runs:
                _set_run_font(run)

            # 追踪当前章节
            if level >= 3:
                sec_key = text[:6].replace(" ", "")
                for candidate in red_sections:
                    if sec_key.startswith(candidate):
                        current_section = candidate
                        break
                else:
                    current_section = ""

                if current_section in red_sections:
                    for run in p.runs:
                        run.font.color.rgb = RED
            continue

        if btype == "code":
            p = doc.add_paragraph()
            run = p.add_run(block["content"])
            _set_run_font(run, font_name=CODE_FONT, font_size=FONT_SIZE_CODE, color=RED)
            continue

        if btype == "paragraph":
            text = block["content"]
            is_red = (current_section in red_sections)

            # 如果当前章节不在red_sections中，再逐段判断是否新增
            if not is_red and len(text) > 15:
                if _is_new_paragraph(text, original_paras):
                    is_red = True

            _write_paragraph(doc, text, is_red)
            continue

    # 保存
    Path(output_docx_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_docx_path)
    print(f"[OK] 修订文档已保存: {output_docx_path}")
    print(f"      原始稿: {original_md_path}")
    print(f"      修订稿: {final_md_path}")
    print(f"      标红章节: {', '.join(sorted(red_sections)) if red_sections else '无'}")


def main():
    parser = argparse.ArgumentParser(description="生成带红色修订标记的论文Word文档")
    parser.add_argument("--original", required=True, help="原始markdown文件路径")
    parser.add_argument("--final", required=True, help="修订后markdown文件路径")
    parser.add_argument("--output", required=True, help="输出docx文件路径")
    parser.add_argument("--title", required=True, help="文档标题")
    parser.add_argument(
        "--red-sections",
        default="",
        help='需要标红的章节编号，逗号分隔，如 "2.2.2,2.2.3,2.3.1"',
    )
    parser.add_argument(
        "--legend",
        default="红色=新增/修改内容；黑色=原文保留。",
        help="修改说明文字",
    )
    args = parser.parse_args()

    red_set = set(s.strip() for s in args.red_sections.split(",") if s.strip())

    generate_revision_docx(
        original_md_path=args.original,
        final_md_path=args.final,
        output_docx_path=args.output,
        title=args.title,
        red_sections=red_set,
        legend_text=args.legend,
    )


if __name__ == "__main__":
    main()
