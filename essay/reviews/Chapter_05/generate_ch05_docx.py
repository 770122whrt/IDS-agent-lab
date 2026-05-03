"""生成第五章 Word 文档"""

import re
from pathlib import Path
from docx import Document
from docx.shared import RGBColor, Pt, Inches, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ── 字体 ──
FONT_CN = "仿宋"
FONT_EN = "Times New Roman"
FONT_CODE = "Courier New"
SIZE_BODY = Pt(12)       # 小四
SIZE_H1 = Pt(16)         # 三号
SIZE_H2 = Pt(14)         # 四号
SIZE_H3 = Pt(12)         # 小四
SIZE_CODE = Pt(10)       # 五号
SIZE_TABLE = Pt(10.5)    # 五号半

RED = RGBColor(255, 0, 0)
BLACK = RGBColor(0, 0, 0)
GREY = RGBColor(100, 100, 100)
BLUE = RGBColor(0, 0, 180)
WHITE = RGBColor(255, 255, 255)
HEADER_BG = "D9E2F3"
NOTE_BG = "FFF2CC"


def set_run(run, font_cn=FONT_CN, size=SIZE_BODY, color=BLACK, bold=False, italic=False):
    run.font.name = font_cn
    run.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run.font.size = size
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic


def add_paragraph_with_bold(doc, text, size=SIZE_BODY, color=BLACK, alignment=None, spacing_after=Pt(6)):
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = spacing_after
    parts = re.split(r"\*\*(.*?)\*\*", text)
    for i, part in enumerate(parts):
        if not part:
            continue
        run = p.add_run(part)
        is_bold = (i % 2 == 1)
        set_run(run, size=size, color=color, bold=is_bold)
    return p


def add_code_block(doc, lines, bg_color="F2F2F2"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): bg_color,
    })
    shading.append(shd)

    for line in lines:
        p = cell.add_paragraph()
        run = p.add_run(line)
        set_run(run, font_cn=FONT_CODE, size=SIZE_CODE, color=BLACK)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    if cell.paragraphs and not cell.paragraphs[0].text:
        p_element = cell.paragraphs[0]._element
        p_element.getparent().remove(p_element)

    doc.add_paragraph()


def add_note_block(doc, lines, bg_color=NOTE_BG):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): bg_color,
    })
    shading.append(shd)

    for line in lines:
        p = cell.add_paragraph()
        parts = re.split(r"\*\*(.*?)\*\*", line)
        for i, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            is_bold = (i % 2 == 1)
            color = BLUE if is_bold else GREY
            set_run(run, size=SIZE_CODE, color=color, bold=is_bold)
        p.paragraph_format.space_after = Pt(2)

    if cell.paragraphs and not cell.paragraphs[0].text:
        p_element = cell.paragraphs[0]._element
        p_element.getparent().remove(p_element)

    doc.add_paragraph()


def add_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for j, header in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run(run, size=SIZE_TABLE, color=WHITE, bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = cell._element.get_or_add_tcPr()
        shd = shading.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "2E75B6",
        })
        shading.append(shd)

    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            cell = table.cell(i + 1, j)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            set_run(run, size=SIZE_TABLE)
            if i % 2 == 1:
                shading = cell._element.get_or_add_tcPr()
                shd = shading.makeelement(qn("w:shd"), {
                    qn("w:val"): "clear",
                    qn("w:color"): "auto",
                    qn("w:fill"): "F2F7FC",
                })
                shading.append(shd)

    if col_widths:
        for i, row_obj in enumerate(table.rows):
            for j, width in enumerate(col_widths):
                row_obj.cells[j].width = width

    doc.add_paragraph()


def parse_and_generate(md_path, docx_path):
    with open(md_path, "r", encoding="utf-8") as f:
        md = f.read()

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = FONT_EN
    style.font.size = SIZE_BODY
    style.element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)

    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    lines = md.split("\n")
    i = 0
    in_code = False
    in_note = False
    code_buffer = []
    note_buffer = []
    para_buffer = []

    def flush_para():
        if para_buffer:
            text = " ".join(para_buffer)
            if text.strip() == "---":
                para_buffer.clear()
                return
            add_paragraph_with_bold(doc, text)
            para_buffer.clear()

    def flush_code():
        if code_buffer:
            add_code_block(doc, code_buffer)
            code_buffer.clear()

    def flush_note():
        if note_buffer:
            add_note_block(doc, note_buffer)
            note_buffer.clear()

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            flush_para()
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(stripped)
            i += 1
            continue

        if stripped.startswith("> **【图表需求") or stripped.startswith("> **【第五章图表"):
            flush_para()
            in_note = True
            note_buffer.append(stripped.lstrip("> "))
            i += 1
            continue

        if in_note:
            if stripped.startswith(">"):
                note_buffer.append(stripped.lstrip("> ").strip())
                i += 1
                continue
            elif stripped == "---":
                flush_note()
                in_note = False
                i += 1
                continue
            else:
                flush_note()
                in_note = False

        if stripped.startswith("#"):
            flush_para()
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped.lstrip("#").strip()
            heading = doc.add_heading(text, level=level)
            for run in heading.runs:
                if level == 1:
                    set_run(run, size=SIZE_H1, bold=True)
                elif level == 2:
                    set_run(run, size=SIZE_H2, bold=True)
                else:
                    set_run(run, size=SIZE_H3, bold=True)
            i += 1
            continue

        if stripped.startswith("|") and "|" in stripped[1:]:
            flush_para()
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row_text = lines[i].strip()
                if re.match(r"^\|[\s\-:]+\|", row_text):
                    i += 1
                    continue
                cells = [c.strip() for c in row_text.split("|")[1:-1]]
                table_lines.append(cells)
                i += 1
            if len(table_lines) >= 2:
                headers = table_lines[0]
                rows = table_lines[1:]
                add_styled_table(doc, headers, rows)
            continue

        if not stripped:
            flush_para()
            i += 1
            continue

        if stripped == "---":
            flush_para()
            i += 1
            continue

        para_buffer.append(stripped)
        i += 1

    flush_para()
    if in_code:
        flush_code()
    if in_note:
        flush_note()

    Path(docx_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)
    print(f"[OK] 文档已保存: {docx_path}")


if __name__ == "__main__":
    md_path = Path(__file__).parent / "CHAPTER_05_FINAL.md"
    docx_path = Path(__file__).parent.parent.parent / "output" / "chapter_05_revised.docx"
    parse_and_generate(str(md_path), str(docx_path))
